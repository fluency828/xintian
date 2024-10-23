import pandas as pd
from docx import Document 
from docx.shared import Cm,Pt
import math
import io
from docx.oxml.ns import qn
# from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

class gen_document():
    def __init__(self,
                 instance,
                #  Large_components_temp_ls,
                #  generator_temp_ls,
                #  pitch_motor_temp_ls,
                 Large_component_fig_ls,
                 Large_components_fig_single_ls,
                 generator_temp_fig_ls,
                 pitch_motor_temp_fig_ls,
                 torque_fig_ls,
                #  yaw_data,
                 blade_pw_fig_ls,
                 blade_time_fig_ls,
                 ):
        self.Large_components_temp_ls = instance.Large_components_temp_ls
        self.Large_components_temp_ls = instance.Large_components_temp_ls
        self.generator_temp_ls = instance.generator_temp_ls
        self.pitch_motor_temp_ls = instance.pitch_motor_temp_ls
        self.Large_component_fig_ls = Large_component_fig_ls
        self.Large_components_fig_single_ls = Large_components_fig_single_ls
        self.generator_temp_fig_ls = generator_temp_fig_ls
        self.pitch_motor_temp_fig_ls = pitch_motor_temp_fig_ls
        self.torque_fig_ls = torque_fig_ls
        # self.yaw_data = yaw_data
        self.blade_pw_fig_ls = blade_pw_fig_ls
        self.blade_time_fig_ls = blade_time_fig_ls
        self.instance = instance
        self.document = Document()
        self.document.styles['Normal'].font.name = 'Times New Roman'
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
        self.set_orientation_landscape()
        self.set_margins()
        self.gen_docx()

    def set_orientation_landscape(self):
        """
        设置页面方向为横向
        """
        section = self.document.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

    def set_margins(self):
        """
        设置页面边距
        """
        section = self.document.sections[0]
        section.top_margin = Cm(3.18)
        section.bottom_margin = Cm(3.18)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    def gen_docx(self):
        # print(1)
        self.document.add_heading('3.1 大部件温度异常',level=2)
        if len(self.Large_component_fig_ls)>0:
            self.gen_Large_component_paragraph()
        # print(1)
        self.document.add_heading('大部件温度温度预警(非满发)',level=3)
        if len(self.Large_components_fig_single_ls)>0:
            self.gen_all_wtg_paragraph(self.Large_components_fig_single_ls)
        else: 
            self.document.add_paragraph('大部件温度温度预警(非满发)无异常')
        
        self.document.add_heading('发电机绕组温度异常的风机',level=3)
        if len(self.generator_temp_fig_ls)>0:
            self.gen_all_wtg_paragraph(self.generator_temp_fig_ls)
        else: 
            self.document.add_paragraph('发电机绕组温度同风机不同相对比无异常')
        # print(1)
        self.document.add_heading('变桨电机温度异常的风机',level=3)
        if len(self.pitch_motor_temp_fig_ls)>0:
            self.gen_all_wtg_paragraph(self.pitch_motor_temp_fig_ls)
        else: 
            self.document.add_paragraph('变桨电机温度同风机不同相对比无异常')
        # print(1)
        self.document.add_heading('3.2 偏航对风',level=2)
        # if self.yaw_data is not None:
        #     self.gen_table_paragraph(self.yaw_data)
        # print(1)
        self.document.add_heading('3.3 转矩控制',level=2)
        if len(self.torque_fig_ls)>0:
            self.gen_all_wtg_paragraph(self.torque_fig_ls)
        # print(1)
        self.document.add_heading('3.4 桨叶角度对零',level=2)
        if len(self.blade_pw_fig_ls)>0:
            self.gen_all_wtg_paragraph(self.blade_pw_fig_ls)
        if len(self.blade_time_fig_ls)>0:
            self.gen_all_wtg_paragraph(self.blade_time_fig_ls)
        
        return self.document
    
    def gen_all_wtg_paragraph(self,figure_list):
        row = math.ceil(len(figure_list)/4)
        table = self.document.add_table(rows=row,cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER # 表格居中对齐
        for i,cell in enumerate(table._cells):
            if i<len(figure_list):
                # run = cell.add_paragraph().add_run()
                buf = io.BytesIO()
                figure_list[i].savefig(buf,dpi=300,facecolor='white',format='jpg',bbox_inches='tight')
                buf.seek(0)
                cell_paragraph = cell.paragraphs[0]
                run = cell_paragraph.add_run()
                run.add_picture(buf,height=Cm(3.2))
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
                buf.close()
            else:
                break

    def gen_Large_component_paragraph(self):
        scene_list = self.instance.Large_components_temp_ls + self.instance.generator_temp_ls
        for i,scene in enumerate(scene_list):
            self.document.add_heading(f'{i+1}、{scene}',level=3)
            table = self.document.add_table(rows=1,cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER # 表格居中对齐
            for j,cell in enumerate(table._cells):
                # run = cell.add_paragraph().add_run()
                # run.add_picture(self.Large_component_fig_ls[i*2+j],height=Cm(5.5))
                buf = io.BytesIO()
                self.Large_component_fig_ls[i*2+j].savefig(buf,dpi=300,facecolor='white',format='jpg',bbox_inches='tight')
                buf.seek(0)
                cell_paragraph = cell.paragraphs[0]
                run = cell_paragraph.add_run()                
                run.add_picture(buf,height=Cm(5.5))
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
                buf.close()

    
    def gen_table_paragraph(self,dataframe):
        table = self.document.add_table(rows=1,cols=dataframe.shape[1])
        table.sytle = 'Table Grid'
        head_cells = table.rows[0].cells
        for i,col in enumerate(dataframe.columns):
            head_cells[i].text = col
        
        for _,row in dataframe.iterrows():
            row_cells = table.add_row().cells
            # print(row)
            for j in range(len(row)):
                # print(j,list(row))
                row_cells[j].text = str(list(row)[j])

